"""Vector store and retriever initialization."""

import os
import tempfile
from pathlib import Path

from functools import partial
from docx import Document
from langchain_core.documents import Document as LangchainDocument
from langchain_community.document_loaders import DirectoryLoader, TextLoader, PyPDFLoader
from langchain_community.vectorstores import SKLearnVectorStore
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from agent.config.settings import (
    VECTOR_STORE_PERSIST_DIR,
    VECTOR_STORE_SERIALIZER,
    DOCS_DIR,
    DOCUMENT_CHUNK_SIZE,
    DOCUMENT_CHUNK_OVERLAP,
    RETRIEVER_LAMBDA_MULT,
)


def load_docx_files(docs_path: Path) -> list:
    """Load all DOCX files from the given path.
    
    Args:
        docs_path: Path to the documents directory
        
    Returns:
        List of LangchainDocument objects
    """
    documents = []
    try:
        docx_paths = list(docs_path.glob("**/*.docx"))
        for docx_path in docx_paths:
            try:
                doc = Document(str(docx_path))
                text = "\n".join([para.text for para in doc.paragraphs])
                if text.strip():
                    documents.append(
                        LangchainDocument(
                            page_content=text,
                            metadata={"source": str(docx_path)}
                        )
                    )
            except Exception as e:
                print(f"Warning: Could not load DOCX file {docx_path}: {e}")
    except Exception as e:
        print(f"Warning: Error processing DOCX files: {e}")
    
    return documents


def get_vector_db_retriever():
    """Initialize or load vector database retriever.
    
    Creates a new vector store from local documentation files (.txt, .pdf, .docx).
    Always rebuilds the vector store from current documents to ensure fresh data.
    
    Returns:
        A retriever instance from SKLearnVectorStore
    """
    persist_path = os.path.join(tempfile.gettempdir(), VECTOR_STORE_PERSIST_DIR)
    embd = OpenAIEmbeddings()

    # Clear old cache on startup to ensure fresh data from current documents
    if os.path.exists(persist_path):
        try:
            import shutil
            shutil.rmtree(persist_path)
            print(f"[INFO] Cleared old vector store cache: {persist_path}")
        except Exception as e:
            print(f"[WARNING] Could not clear vector store cache: {e}")

    # Otherwise, index documents from local directory and create new vector store
    # Resolve path relative to the project root
    docs_path = Path(__file__).parent.parent.parent.parent / DOCS_DIR
    if not docs_path.exists():
        raise ValueError(f"Documentation directory not found: {docs_path}")
    
    documents = []
    
    # Load .txt files
    text_loader_utf8 = partial(TextLoader, encoding="utf-8")
    text_loader = DirectoryLoader(
        path=str(docs_path),
        glob="**/*.txt",
        loader_cls=text_loader_utf8,
        show_progress=True,
        use_multithreading=True
    )
    documents.extend(text_loader.load())
    
    # Load .pdf files
    pdf_loader = DirectoryLoader(
        path=str(docs_path),
        glob="**/*.pdf",
        loader_cls=PyPDFLoader,
        show_progress=True,
        use_multithreading=True
    )
    try:
        documents.extend(pdf_loader.load())
    except Exception as e:
        print(f"Warning: Could not load PDF files: {e}")
    
    # Load .docx files
    documents.extend(load_docx_files(docs_path))
    
    if not documents:
        raise ValueError(f"No documents found in {docs_path}")

    text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=DOCUMENT_CHUNK_SIZE,
        chunk_overlap=DOCUMENT_CHUNK_OVERLAP
    )
    doc_splits = text_splitter.split_documents(documents)

    vectorstore = SKLearnVectorStore.from_documents(
        documents=doc_splits,
        embedding=embd,
        persist_path=persist_path,
        serializer=VECTOR_STORE_SERIALIZER
    )
    vectorstore.persist()
    return vectorstore.as_retriever(lambda_mult=RETRIEVER_LAMBDA_MULT)
